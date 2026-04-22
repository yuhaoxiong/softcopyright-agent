"""DOCX writer using python-docx for professional soft copyright documents.

Generates Word documents with:
- Cover page (title, version, date)
- Chinese fonts (SimSun body, SimHei headings)
- Proper table rendering from Markdown
- Header (title) and footer (page number)
- 1.5x line spacing, A4 page, standard margins
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ── 字体常量 ─────────────────────────────────────────────────────
_FONT_BODY = "宋体"
_FONT_HEADING = "黑体"
_FONT_BODY_EN = "Times New Roman"
_FONT_SIZE_BODY = Pt(12)        # 小四
_FONT_SIZE_H1 = Pt(16)          # 三号
_FONT_SIZE_H2 = Pt(14)          # 四号
_FONT_SIZE_H3 = Pt(12)          # 小四加粗
_FONT_SIZE_COVER_TITLE = Pt(22) # 二号
_LINE_SPACING = 1.5


def write_docx(
    path: Path,
    title: str,
    markdown_text: str,
    *,
    enable_remote_diagrams: bool = False,
) -> None:
    """Write a professionally formatted DOCX from Markdown-like text."""

    doc = Document()

    _setup_default_font(doc)
    _setup_page(doc)
    _add_cover_page(doc, title)
    _add_header_footer(doc, title)
    _render_markdown(doc, markdown_text, enable_remote_diagrams=enable_remote_diagrams)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


# ── 文档基础设置 ──────────────────────────────────────────────────

def _setup_default_font(doc: Document) -> None:
    """Set default font for the entire document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = _FONT_BODY_EN
    font.size = _FONT_SIZE_BODY
    # 设置中文字体
    style.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_BODY)

    pf = style.paragraph_format
    pf.line_spacing = _LINE_SPACING
    pf.space_after = Pt(4)
    pf.space_before = Pt(2)

    # 设置标题样式
    for level, (size, heading_name) in enumerate(
        [(Pt(16), "Heading 1"), (Pt(14), "Heading 2"), (Pt(12), "Heading 3")],
        start=1,
    ):
        try:
            h_style = doc.styles[heading_name]
        except KeyError:
            continue
        h_font = h_style.font
        h_font.name = _FONT_BODY_EN
        h_font.size = size
        h_font.bold = True
        h_font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        h_style.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_HEADING)
        h_pf = h_style.paragraph_format
        h_pf.space_before = Pt(12)
        h_pf.space_after = Pt(6)
        h_pf.line_spacing = _LINE_SPACING


def _setup_page(doc: Document) -> None:
    """Configure A4 page with standard margins."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


# ── 封面页 ────────────────────────────────────────────────────────

def _add_cover_page(doc: Document, title: str) -> None:
    """Add a cover page with title, version, and date."""
    # 上方空白
    for _ in range(6):
        doc.add_paragraph("")

    # 标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.font.size = _FONT_SIZE_COVER_TITLE
    run.font.bold = True
    run.font.name = _FONT_BODY_EN
    run.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_HEADING)

    # 副标题
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("软件设计说明书")
    run.font.size = Pt(18)
    run.font.name = _FONT_BODY_EN
    run.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_HEADING)

    doc.add_paragraph("")

    # 版本和日期信息
    info_lines = [
        f"版本号：V1.0",
        f"编制日期：{date.today().strftime('%Y年%m月%d日')}",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(14)
        run.font.name = _FONT_BODY_EN
        run.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_BODY)

    # 分页
    doc.add_page_break()


# ── 页眉页脚 ──────────────────────────────────────────────────────

def _add_header_footer(doc: Document, title: str) -> None:
    """Add header with title and footer with page number."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # 页眉（非首页）
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(9)
    run.font.name = _FONT_BODY_EN
    run.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_BODY)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 页脚页码（非首页）
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(p)


def _add_page_number_field(paragraph) -> None:
    """Insert a PAGE field code for auto page numbering."""
    from docx.oxml import OxmlElement

    run = paragraph.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run.element.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run.element.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run.element.append(fld_char_end)


# ── Markdown 渲染 ─────────────────────────────────────────────────

def _render_markdown(doc: Document, markdown_text: str, *, enable_remote_diagrams: bool = False) -> None:
    """Parse Markdown text and render into the document."""
    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行跳过
        if not stripped:
            i += 1
            continue

        # 标题
        if stripped.startswith("### "):
            _add_heading(doc, stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            _add_heading(doc, stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            _add_heading(doc, stripped[2:].strip(), level=1)
            i += 1
            continue

        # Mermaid 图表
        if stripped.startswith("```mermaid"):
            mermaid_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                mermaid_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # 跨过结束标签 ```
            
            mermaid_text = "\n".join(mermaid_lines)
            try:
                import io
                from .diagram_renderer import render_mermaid_to_png
                png_data = render_mermaid_to_png(mermaid_text, allow_remote=enable_remote_diagrams)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                # 配合 A4 页面留白，宽15cm左右比较合适
                run.add_picture(io.BytesIO(png_data), width=Cm(15))
            except Exception as e:
                if "disabled" not in str(e).lower():
                    print(f"渲染 Mermaid 失败: {e}")
                # 失败则回退写入源码
                for line in mermaid_lines:
                    p = doc.add_paragraph()
                    _set_run_font(p.add_run(line))
            continue

        # Markdown 代码块的回退忽略
        if stripped.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            if i < len(lines):
                i += 1
            continue

        # Markdown 表格
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            continue

        # 列表项
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _set_run_font(p.add_run(text))
            i += 1
            continue

        # 有序列表
        num_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if num_match:
            text = num_match.group(2).strip()
            p = doc.add_paragraph(style="List Number")
            _set_run_font(p.add_run(text))
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _set_run_font(p.add_run(stripped.lstrip("- ").strip()))
        i += 1


def _add_heading(doc: Document, text: str, level: int) -> None:
    """Add a heading with proper Chinese font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = _FONT_BODY_EN
        run.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_HEADING)


def _add_table(doc: Document, table_lines: list[str]) -> None:
    """Parse Markdown table lines and add a Word table."""
    if len(table_lines) < 2:
        # 不够组成表格，当普通段落处理
        for line in table_lines:
            p = doc.add_paragraph()
            _set_run_font(p.add_run(line))
        return

    # 解析表头
    header_cells = _parse_table_row(table_lines[0])

    # 跳过分隔行（|---|---|）
    data_start = 1
    if data_start < len(table_lines) and re.match(r"^\|[\s\-:|]+\|$", table_lines[data_start]):
        data_start = 2

    # 解析数据行
    data_rows = [_parse_table_row(line) for line in table_lines[data_start:]]
    if not header_cells:
        return

    col_count = len(header_cells)
    row_count = 1 + len(data_rows)  # 表头 + 数据行

    table = doc.add_table(rows=row_count, cols=col_count)
    table.style = "Table Grid"

    # 写入表头
    for col_idx, cell_text in enumerate(header_cells):
        if col_idx < col_count:
            cell = table.rows[0].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.bold = True
            run.font.size = Pt(10.5)
            run.font.name = _FONT_BODY_EN
            run.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_BODY)

    # 写入数据行
    for row_idx, row_cells in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row_cells):
            if col_idx < col_count:
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.size = Pt(10.5)
                run.font.name = _FONT_BODY_EN
                run.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_BODY)

    # 表格后空一行
    doc.add_paragraph("")


def _parse_table_row(line: str) -> list[str]:
    """Split a Markdown table row into cell texts."""
    # 去掉首尾的 |
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _set_run_font(run) -> None:
    """Apply the standard body font to a run."""
    run.font.size = _FONT_SIZE_BODY
    run.font.name = _FONT_BODY_EN
    run.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_BODY)
